from flask import Response,render_template, request, redirect, url_for, flash, jsonify
from modules.fournisseur import fournisseur_bp
from models.db import Fournisseur, Session
from modules.fournisseur.repository.fournisseur import FournisseurRepository
from modules.fournisseur.service.fournisseur import FournisseurService
from sqlalchemy import create_engine
import os
from werkzeug.utils import secure_filename
from io import BytesIO
from fpdf import FPDF
import pandas as pd
from docx import Document

# Create a session instance
db_session = Session()

# Repository and Service
fournisseur_repository = FournisseurRepository(db_session)
fournisseur_service = FournisseurService(fournisseur_repository)

UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

# Ensure the upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

VALID_ETIQUETTES = {'Fournisseur_Produit', 'Service_Transit', 'Fournisseur_Fonctionnement'}

def validate_etiquettes(etiquettes):
    if not all(etiquette in VALID_ETIQUETTES for etiquette in etiquettes):
        raise ValueError(f"Invalid etiquettes: {etiquettes}. Must be one of {VALID_ETIQUETTES}.")
    
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@fournisseur_bp.route('/export/pdf', methods=['GET'])
def export_pdf():
    # Get search and filter parameters
    search_query = request.args.get('search', '').strip()
    filter_trader = request.args.get('filter_trader') == 'on'
    filter_list_noire = request.args.get('filter_list_noire') == 'on'

    # Apply filters
    query = db_session.query(Fournisseur)
    if search_query:
        query = query.filter(Fournisseur.nom.ilike(f'%{search_query}%'))
    if filter_trader:
        query = query.filter(Fournisseur.trader.is_(True))
    if filter_list_noire:
        query = query.filter(Fournisseur.list_noire.is_(True))
    fournisseurs = query.all()

    # Create a PDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Title
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, txt="Liste des Fournisseurs", ln=True, align="C")
    pdf.ln(10)

    # Table Header
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(40, 10, "Nom", border=1)
    pdf.cell(40, 10, "Numero", border=1)
    pdf.cell(60, 10, "Email", border=1)
    pdf.cell(30, 10, "Trader", border=1)
    pdf.cell(30, 10, "Liste Noire", border=1)
    pdf.ln()

    # Table Rows
    pdf.set_font("Arial", size=12)
    for fournisseur in fournisseurs:
        pdf.cell(40, 10, fournisseur.nom, border=1)
        pdf.cell(40, 10, fournisseur.numero, border=1)
        pdf.cell(60, 10, fournisseur.email, border=1)
        pdf.cell(30, 10, "Oui" if fournisseur.trader else "Non", border=1)
        pdf.cell(30, 10, "Oui" if fournisseur.list_noire else "Non", border=1)
        pdf.ln()

    # Return PDF as response
    response = Response(pdf.output(dest='S').encode('latin1'))
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'inline; filename=fournisseurs.pdf'
    return response
@fournisseur_bp.route('/export/excel', methods=['GET'])
def export_excel():
    # Get search and filter parameters
    search_query = request.args.get('search', '').strip()
    filter_trader = request.args.get('filter_trader') == 'on'
    filter_list_noire = request.args.get('filter_list_noire') == 'on'

    # Apply filters
    query = db_session.query(Fournisseur)
    if search_query:
        query = query.filter(Fournisseur.nom.ilike(f'%{search_query}%'))
    if filter_trader:
        query = query.filter(Fournisseur.trader.is_(True))
    if filter_list_noire:
        query = query.filter(Fournisseur.list_noire.is_(True))
    fournisseurs = query.all()

    # Create a DataFrame
    data = [{
        "Nom": fournisseur.nom,
        "Numero": fournisseur.numero,
        "Email": fournisseur.email,
        "Trader": "Oui" if fournisseur.trader else "Non",
        "Liste Noire": "Oui" if fournisseur.list_noire else "Non"
    } for fournisseur in fournisseurs]
    df = pd.DataFrame(data)

    # Save to Excel
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Fournisseurs')

    # Return Excel as response
    response = Response(output.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response.headers['Content-Disposition'] = 'attachment; filename=fournisseurs.xlsx'
    return response

@fournisseur_bp.route('/export/word', methods=['GET'])
def export_word():
    # Get search and filter parameters
    search_query = request.args.get('search', '').strip()
    filter_trader = request.args.get('filter_trader') == 'on'
    filter_list_noire = request.args.get('filter_list_noire') == 'on'

    # Apply filters
    query = db_session.query(Fournisseur)
    if search_query:
        query = query.filter(Fournisseur.nom.ilike(f'%{search_query}%'))
    if filter_trader:
        query = query.filter(Fournisseur.trader.is_(True))
    if filter_list_noire:
        query = query.filter(Fournisseur.list_noire.is_(True))
    fournisseurs = query.all()

    # Create a Word Document
    doc = Document()
    doc.add_heading('Liste des Fournisseurs', level=1)

    # Add Table
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nom'
    hdr_cells[1].text = 'Numero'
    hdr_cells[2].text = 'Email'
    hdr_cells[3].text = 'Trader'
    hdr_cells[4].text = 'Liste Noire'

    for fournisseur in fournisseurs:
        row_cells = table.add_row().cells
        row_cells[0].text = fournisseur.nom
        row_cells[1].text = fournisseur.numero
        row_cells[2].text = fournisseur.email
        row_cells[3].text = "Oui" if fournisseur.trader else "Non"
        row_cells[4].text = "Oui" if fournisseur.list_noire else "Non"

    # Save to Word
    output = BytesIO()
    doc.save(output)

    # Return Word as response
    response = Response(output.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers['Content-Disposition'] = 'attachment; filename=fournisseurs.docx'
    return response

@fournisseur_bp.route('/', methods=['GET'])
def list_fournisseurs():
    search_query = request.args.get('search', '').strip()
    filter_trader = request.args.get('filter_trader') == 'on'
    filter_list_noire = request.args.get('filter_list_noire') == 'on'
    filter_etiquette = request.args.get('filter_etiquette', '').strip()

    query = db_session.query(Fournisseur)

    # Apply search filter if search_query is not empty
    if search_query:
        query = query.filter(Fournisseur.nom.ilike(f'%{search_query}%'))

    # Apply Trader filter if checkbox is selected
    if filter_trader:
        query = query.filter(Fournisseur.trader.is_(True))

    # Apply Liste Noire filter if checkbox is selected
    if filter_list_noire:
        query = query.filter(Fournisseur.list_noire.is_(True))

    # Apply Etiquettes filter if provided
    if filter_etiquette:
        query = query.filter(Fournisseur.etiquettes.contains([filter_etiquette]))

    fournisseurs = query.all()

    return render_template('fournisseurs/list_fournisseurs.html', fournisseurs=fournisseurs, valid_etiquettes=VALID_ETIQUETTES)
@fournisseur_bp.route('/ajouter', methods=['GET', 'POST'])
def ajouter_fournisseur():
    if request.method == 'POST':
        # Get form data
        nom = request.form.get('nom')
        NIF = request.form.get('NIF')
        contact = request.form.get('contact')
        address = request.form.get('address')
        email = request.form.get('email')
        image_path = f"uploads/default.jpg"
        trader = request.form.get('trader') == 'on'  # Checkbox value
        list_noire = request.form.get('list_noire') == 'on'  # Checkbox value

        # Handle file upload
        if 'sigle_image' in request.files:
            file = request.files['sigle_image']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(UPLOAD_FOLDER, filename))  # Save the file to the upload folder
                image_path = f"uploads/{filename}"  # Save the relative path (e.g., "uploads/image.jpg")

        # Handle etiquettes (multi-select or checkboxes)
        etiquettes = request.form.getlist('etiquettes')  # Get selected values as a list
        validate_etiquettes(etiquettes)  # Validate the input

        # Create a new Fournisseur object
        new_fournisseur = Fournisseur(
            nom=nom,
            nif=nif,
            contact=contact,
            address=address,
            email=email,
            sigle_image=image_path,
            trader=trader,
            list_noire=list_noire,
            etiquettes=etiquettes  # Save as JSON or a comma-separated string
        )

        # Add to the database
        try:
            db_session.add(new_fournisseur)
            db_session.commit()
            flash('Fournisseur ajouté avec succès!', 'success')
            return redirect(url_for('fournisseurs.list_fournisseurs'))
        except Exception as e:
            db_session.rollback()
            flash(f'Erreur: {str(e)}', 'danger')

    return render_template('fournisseurs/ajouter_fournisseur.html', valid_etiquettes=VALID_ETIQUETTES)

@fournisseur_bp.route('/<int:id>/delete', methods=['POST'])
def delete_fournisseur(id):
    fournisseur = db_session.query(Fournisseur).get(id)
    if not fournisseur:
        flash('Fournisseur introuvable.', 'error')
        return redirect(url_for('fournisseurs.list_fournisseurs'))
    try:
        db_session.delete(fournisseur)
        db_session.commit()
        flash('Fournisseur supprimé avec succès.', 'success')
    except Exception as e:
        db_session.rollback()
        flash(f'Erreur lors de la suppression: {str(e)}', 'danger')
    return redirect(url_for('fournisseurs.list_fournisseurs'))

@fournisseur_bp.route('/fournisseurs/<int:id>/edit', methods=['GET', 'POST'])
def edit_fournisseur(id):
    fournisseur = db_session.query(Fournisseur).get(id)
    if not fournisseur:
        flash('Fournisseur introuvable.', 'error')
        return redirect(url_for('fournisseurs.list_fournisseurs'))

    if request.method == 'POST':
        data = request.form
        fournisseur.nom = data.get('nom')
        fournisseur.NIF = data.get('NIF')
        print("Updated NIF:", fournisseur.NIF) 
        fournisseur.contact = data.get('contact')
        fournisseur.address = data.get('address')
        fournisseur.email = data.get('email')
        fournisseur.trader = request.form.get('trader') == 'on'  # Checkbox value
        fournisseur.list_noire = request.form.get('list_noire') == 'on'  # Checkbox value

        # Handle file upload for sigle_image
        if 'sigle_image' in request.files:
            file = request.files['sigle_image']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(UPLOAD_FOLDER, filename))  # Save the file to the upload folder
                fournisseur.sigle_image = f"uploads/{filename}"  # Save the relative path (e.g., "uploads/image.jpg")

        # Handle etiquettes (multi-select or checkboxes)
        etiquettes = request.form.getlist('etiquettes')  # Get selected values as a list
        validate_etiquettes(etiquettes)  # Validate the input
        fournisseur.etiquettes = etiquettes  # Save as JSON or a comma-separated string

        try:
            db_session.commit()
            flash('Fournisseur modifié avec succès.', 'success')
            return redirect(url_for('fournisseurs.list_fournisseurs'))
        except Exception as e:
            db_session.rollback()
            flash(f'Erreur lors de la mise à jour: {str(e)}', 'danger')

    return render_template('fournisseurs/edit_fournisseur.html', fournisseur=fournisseur, valid_etiquettes=VALID_ETIQUETTES)


@fournisseur_bp.route('/autocomplete', methods=['GET'])
def autocomplete_fournisseurs():
    query = request.args.get('query', '').strip()
    if not query:
        return jsonify([])

    # Search fournisseurs by name (case-insensitive)
    fournisseurs = db_session.query(Fournisseur).filter(Fournisseur.nom.ilike(f'%{query}%')).all()
    results = [{'id': fournisseur.id, 'nom': fournisseur.nom} for fournisseur in fournisseurs]

    return jsonify(results)